from src.document_generator.base_document_generator import BaseDocumentGenerator
from faker import Faker
import datetime
from docxtpl import DocxTemplate
import tempfile
from pathlib import Path
import random
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from pdf2image import convert_from_path



fake = Faker("ru_RU")

class InterestFreeLoanAgreement(BaseDocumentGenerator):
    def __init__(self):
        super().__init__()
        self._template_path = Path("src/templates/multipage_docs/test.docx")
        self._doc_type = "interest_free_loan_agreement"
    def generate_contract_number(self):
    # Получаем текущую дату в формате ГГГГММДД
        today = datetime.date.today().strftime("%Y%m%d")
        
        # Генерируем случайный код (например, из 6 цифр)
        random_code = random.randint(100000, 999999)
        
        # Собираем номер договора
        contract_number = f"CON-{today}-{random_code}"
        
        return contract_number
    
    def pdf_to_images(self, pdf_file):
        """
        Convert a .pdf file to images using pdf2image.
        """
        print(f"CONVERT TO PDF PATH: {pdf_file}")
        images = convert_from_path(pdf_file)
        return images

    
    def docx_to_pdf(self, docx_path, pdf_path):
        print(pdf_path)
        pdfmetrics.registerFont(TTFont("DejaVuSans", "./src/Arial/arial.ttf"))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", "./src/Arial/arial_bold.ttf"))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Italic", "./src/Arial/arial_italic.ttf"))

        # Функция для определения выравнивания
        def get_alignment(alignment):
            """Конвертирует выравнивание из python-docx в ReportLab."""
            if alignment == 0:  
                return TA_LEFT
            elif alignment == 1:  
                return TA_CENTER
            elif alignment == 2: 
                return TA_RIGHT
            elif alignment == 3:  
                return TA_JUSTIFY
            else:
                return TA_LEFT  

        def _docx_to_pdf(docx_path, pdf_path):
            """Конвертирует .docx в .pdf с поддержкой текста, выравнивания и таблиц."""
            # Загружаем .docx файл
            document = Document(docx_path)
            
            # Настраиваем PDF
            pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
            story = []  # Список элементов для PDF
            
            # Настройка базового стиля
            styles = getSampleStyleSheet()
            base_style = ParagraphStyle(
                'Base',
                parent=styles['Normal'],
                fontName='DejaVuSans',  # Используем зарегистрированный шрифт
                fontSize=12,
                leading=14
            )
            
            # Обходим содержимое документа
            for para in document.paragraphs:
                # Создаем строку для текста с проверкой каждого фрагмента (run)
                text = ''
                for run in para.runs:
                    # Проверяем жирность и курсив в тексте
                    if run.bold:
                        text += f"<b>{run.text}</b>"  # HTML-теги для жирного текста
                    elif run.italic:
                        text += f"<i>{run.text}</i>"  # HTML-теги для курсива
                    else:
                        text += run.text
                
                if text.strip():  # Пропускаем пустые абзацы
                    # Копируем базовый стиль и добавляем выравнивание
                    paragraph_style = base_style.clone('ParagraphStyle')
                    paragraph_style.alignment = get_alignment(para.alignment)
                    
                    # Обработка жирного шрифта
                    if '<b>' in text:  # Если есть жирный текст
                        paragraph_style.fontName = 'DejaVuSans-Bold'  # Используем жирный шрифт
                    
                    # Добавляем текст как параграф с выравниванием
                    story.append(Paragraph(text, paragraph_style))
                    story.append(Spacer(1, 12))  # Добавляем отступ между абзацами

            # Обработка таблиц
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Извлекаем текст из ячеек
                        cell_text = ''
                        for run in cell.paragraphs[0].runs:
                            if run.bold:
                                cell_text += f"<b>{run.text}</b>"
                            elif run.italic:
                                cell_text += f"<i>{run.text}</i>"
                            else:
                                cell_text += run.text
                        
                        # Добавляем переносы строк с использованием <br/>
                        cell_text = cell_text.replace("\n", "<br/>")  # Заменяем переносы строк на HTML тег <br/>

                        # Создаем объект Paragraph с автоматическим переносом
                        paragraph_style = getSampleStyleSheet()['Normal']
                        paragraph_style.fontName = 'DejaVuSans'  # Шрифт для текста

                        # Используем Paragraph для ячейки, чтобы обеспечить перенос текста
                        paragraph = Paragraph(cell_text.strip(), paragraph_style)
                        row_data.append(paragraph)  # Добавляем Paragraph в строку
                    table_data.append(row_data)
                
                # Создаем таблицу для PDF с автоматическим переносом текста
                pdf_table = Table(table_data)

                # Настроим стиль таблицы
                pdf_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Выравнивание текста в ячейках
                    ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans'),  # Шрифт для текста
                    ('FONTSIZE', (0, 0), (-1, -1), 9),  # Размер шрифта
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Отступы
                    ('WORDSPACE', (0, 0), (-1, -1), 2),  # Добавление интервала между словами
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Вертикальное выравнивание текста в ячейках
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),  # Добавляем отступы внутри ячеек
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('WORDWRAP', (0, 0), (-1, -1), 1),  # Включаем перенос текста в ячейках
                ]))
                # Добавляем таблицу в PDF
                story.append(pdf_table)
                story.append(Spacer(1, 12))  
            pdf.build(story)
            print(f"PDF сохранен в {pdf_path}")
        _docx_to_pdf(docx_path, pdf_path)

    def _generate_one_sample(self):
        context = {
            "org_name": fake.large_company(),
            "org_reason": random.choice(['устава компании', 'договоренности', 'решения учредителей', 'договора', 'протокола собрания']),
            "ip_name": " ".join(fake.passport_owner()),
            "orgnip": fake.individuals_ogrn(),
            "money": round(random.uniform(0, 600000), 2),
            "date": fake.date(),
            "debt_money": round(random.uniform(0, 600000), 2),
            "org_address": fake.address(),
            "org_ogrn": fake.businesses_ogrn(),
            "org_tins": fake.businesses_inn(),
            "org_kpp": fake.kpp(),
            "org_current_account": fake.checking_account(),
            "org_bank_name": fake.bank(),
            "org_correspondent_account_number": fake.correspondent_account(),
            "org_bik": fake.bic(),
            "ip_passport": fake.passport_number(),
            "ip_passport_place":  f"Управление МВД России по г. {fake.city_name()}, {fake.region()}",
            "ip_address": fake.address(),
            "ip_tins": fake.individuals_inn(),
            "ip_ogrnip": fake.individuals_ogrn(),
            "ip_current_account": fake.checking_account(),
            "ip_bank_name": fake.bank(),
            "ip_correspondent_account_number": fake.correspondent_account(),
            "ip_bik": fake.bic(),
            "contract_number": self.generate_contract_number()
        }
        # Загружаем шаблон DOCX
        doc = DocxTemplate(self._template_path)
        # Заменяем метки на данные
        doc.render(context)
        # Сохраняем новый документ в промежуточный файл
        output_path = tempfile.NamedTemporaryFile(dir='./temp', suffix='.docx')
        doc.save(output_path.name)

        pdf_path = tempfile.NamedTemporaryFile(dir='./temp', suffix='.pdf')
        self.docx_to_pdf(output_path.name, pdf_path.name)
        images = self.pdf_to_images(pdf_path.name)
        pdf_path.close()
        output_path.close()
        return images, ""
        

if __name__=="__main__":
    itgfgf = InterestFreeLoanAgreement()
    output_path = itgfgf.generate(100)
    




